from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT_PATH = "docs/BA_BQDrive_Tai_lieu_nghiep_vu.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header, bold=True)
        set_cell_shading(header_cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value))
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return paragraph


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FFF2CC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)
    p.add_run("\n" + body)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True


def build_document():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TÀI LIỆU BA HỆ THỐNG BQDRIVE")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Business Analysis Document - Website đặt thuê xe trực tuyến")
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Phiên bản: 1.0 | Ngày lập: 15/07/2026 | Nguồn rà soát: Source code BQDrive hiện tại").bold = True

    doc.add_paragraph()
    add_note(
        doc,
        "Ghi chú phạm vi",
        "Tài liệu này tổng hợp theo source code hiện tại của BQDrive. Các nghiệp vụ chưa có trong code như bảo trì xe, BANKING/manual, DIRECT, phí phát sinh sau chuyến thuê hoặc luồng hoàn tiền đầy đủ được đưa vào mục giới hạn/hướng phát triển, không xem là chức năng đã hoàn chỉnh.",
    )

    add_heading(doc, "1. Mục đích tài liệu", 1)
    doc.add_paragraph(
        "Tài liệu BA mô tả tổng quan hệ thống BQDrive, phạm vi chức năng, vai trò người dùng, luồng nghiệp vụ, trạng thái dữ liệu, quy tắc xử lý và các điểm tích hợp chính. Tài liệu dùng làm cơ sở cho báo cáo luận văn, kiểm thử nghiệp vụ và phát triển tiếp hệ thống."
    )

    add_heading(doc, "2. Tổng quan hệ thống", 1)
    doc.add_paragraph(
        "BQDrive là hệ thống website hỗ trợ tìm kiếm, đặt thuê và quản lý xe cho thuê. Hệ thống phục vụ ba nhóm người dùng chính: người thuê xe, chủ xe doanh nghiệp, người dùng ký gửi xe và quản trị viên. Các nghiệp vụ trọng tâm gồm quản lý tài khoản, quản lý xe, duyệt xe, tìm kiếm xe, giữ xe tạm thời, tạo booking, duyệt booking, thanh toán, hợp đồng, bàn giao, hoàn tất chuyến thuê, bản đồ vị trí xe và thống kê."
    )

    add_heading(doc, "3. Công nghệ sử dụng", 1)
    add_table(
        doc,
        ["Nhóm", "Công nghệ", "Vai trò trong hệ thống"],
        [
            ["Frontend", "ReactJS, TypeScript, Vite", "Xây dựng giao diện web, routing và tương tác người dùng."],
            ["Giao diện", "Tailwind CSS, Lucide React, React Hot Toast", "Thiết kế giao diện, icon và thông báo thao tác."],
            ["Backend", "Node.js, ExpressJS, TypeScript", "Xây dựng RESTful API và xử lý nghiệp vụ."],
            ["Database", "MongoDB, Mongoose", "Lưu trữ dữ liệu người dùng, xe, booking, payment, contract."],
            ["Xác thực", "JWT, bcryptjs, Google OAuth", "Đăng nhập, phân quyền, mã hóa mật khẩu, đăng nhập Google."],
            ["Email", "Nodemailer/SMTP", "Gửi OTP, quên mật khẩu và email thông báo nghiệp vụ."],
            ["Upload", "Multer/base64 image", "Xử lý hình ảnh xe và giấy tờ người thuê."],
            ["Thanh toán", "MoMo, VNPAY, CASH", "Thanh toán cọc, toàn bộ hoặc phần còn lại."],
            ["Bản đồ", "Leaflet, react-leaflet, geocode/route API", "Chọn vị trí nhận xe, hiển thị bản đồ và tìm đường."],
        ],
        [1.2, 2.2, 3.8],
    )

    add_heading(doc, "4. Tác nhân và phân quyền", 1)
    add_table(
        doc,
        ["Tác nhân", "Mô tả", "Quyền chính"],
        [
            ["USER", "Người dùng thông thường.", "Đăng ký, đăng nhập, tìm kiếm xe, thêm giỏ, đặt xe, thanh toán, xem booking, hợp đồng và lịch sử thanh toán."],
            ["USER ký gửi", "USER có đăng xe ký gửi cho thuê.", "Quản lý xe ký gửi, xem booking phát sinh từ xe ký gửi, xử lý booking, quản lý vị trí xe và thanh toán liên quan."],
            ["BUSINESS", "Chủ xe doanh nghiệp do Admin cấp hoặc được duyệt.", "Quản lý hồ sơ doanh nghiệp, xe doanh nghiệp, booking, bàn giao, hoàn tất chuyến thuê, thanh toán và doanh thu."],
            ["ADMIN", "Quản trị viên hệ thống.", "Quản lý người dùng, doanh nghiệp, thương hiệu, xe, ngày lễ, duyệt xe, khóa/mở khóa/xóa mềm tài khoản."],
        ],
        [1.3, 2.2, 4.0],
    )

    add_heading(doc, "5. Phạm vi chức năng", 1)
    add_heading(doc, "5.1. Chức năng chung", 2)
    add_bullets(
        doc,
        [
            "Đăng ký tài khoản USER bằng email và OTP.",
            "Đăng nhập bằng email/mật khẩu hoặc Google OAuth.",
            "Quên mật khẩu, gửi OTP đặt lại mật khẩu, xác minh OTP và đổi mật khẩu.",
            "Xác thực API bằng JWT Bearer Token.",
            "Phân quyền theo role ADMIN, BUSINESS và USER.",
            "Xem xe công khai, tìm kiếm/lọc xe và xem chi tiết xe.",
        ],
    )

    add_heading(doc, "5.2. Chức năng của ADMIN", 2)
    add_bullets(
        doc,
        [
            "Xem dashboard tổng quan số lượng người dùng, doanh nghiệp, xe, booking và doanh thu.",
            "Quản lý user: xem danh sách, lọc, khóa, mở khóa, xóa mềm.",
            "Cấp tài khoản BUSINESS qua quy trình OTP và tạo doanh nghiệp.",
            "Quản lý doanh nghiệp: xem danh sách, khóa, mở khóa, xóa mềm.",
            "Quản lý thương hiệu xe: thêm, cập nhật, xóa mềm, xem danh sách.",
            "Quản lý xe toàn hệ thống: xem danh sách, lọc trạng thái, duyệt xe, từ chối xe kèm lý do.",
            "Quản lý ngày lễ: thêm, sửa, xóa mềm, bật/tắt áp dụng ngày lễ.",
            "Xem thông báo tổng hợp như xe chờ duyệt, doanh nghiệp cần xử lý và dữ liệu cần kiểm tra.",
        ],
    )

    add_heading(doc, "5.3. Chức năng của BUSINESS", 2)
    add_bullets(
        doc,
        [
            "Xem dashboard doanh nghiệp: số xe, xe chờ duyệt, xe đã duyệt, booking, doanh thu.",
            "Cập nhật hồ sơ doanh nghiệp.",
            "Thêm, cập nhật, xóa mềm, ẩn/hiện xe thuộc doanh nghiệp.",
            "Quản lý vị trí nhận xe trên bản đồ, kéo marker để cập nhật vị trí.",
            "Xem danh sách booking của xe doanh nghiệp.",
            "Duyệt hoặc từ chối booking, nhập lý do từ chối.",
            "Bàn giao xe, xác nhận thu tiền mặt phần còn lại, hoàn tất chuyến thuê hoặc xử lý no-show.",
            "Xem lịch sử booking và lịch sử thanh toán của doanh nghiệp.",
        ],
    )

    add_heading(doc, "5.4. Chức năng của USER thuê xe", 2)
    add_bullets(
        doc,
        [
            "Tìm kiếm xe theo địa điểm, thương hiệu, loại xe, số ghế, giá, nhiên liệu, hộp số và thời gian thuê.",
            "Xem chi tiết xe, hình ảnh, thông tin kỹ thuật, giá thuê và bản đồ vị trí nhận xe.",
            "Tìm đường đến vị trí nhận xe bằng vị trí hiện tại hoặc điểm bắt đầu tự chọn.",
            "Thêm xe vào giỏ để giữ xe tạm thời.",
            "Tạo booking trực tiếp từ trang chi tiết xe hoặc từ giỏ xe.",
            "Nhập thông tin người thuê, CCCD, giấy phép lái xe và ghi chú.",
            "Theo dõi trạng thái booking, thanh toán, hợp đồng và lịch sử thanh toán.",
            "Thanh toán cọc, thanh toán toàn bộ hoặc thanh toán phần còn lại qua MoMo, VNPAY hoặc tiền mặt.",
        ],
    )

    add_heading(doc, "5.5. Chức năng của USER ký gửi xe", 2)
    add_bullets(
        doc,
        [
            "Truy cập khu vực ký gửi xe tại /consignment.",
            "Thêm, cập nhật, xóa mềm, ẩn/hiện xe ký gửi của chính mình.",
            "Theo dõi trạng thái duyệt xe ký gửi.",
            "Quản lý vị trí xe ký gửi trên bản đồ.",
            "Xem và xử lý booking phát sinh từ xe ký gửi.",
            "Xem lịch sử booking và thanh toán liên quan đến xe ký gửi.",
        ],
    )

    add_heading(doc, "6. Trạng thái và danh mục dữ liệu", 1)
    add_table(
        doc,
        ["Nhóm", "Giá trị", "Ý nghĩa"],
        [
            ["CarStatus", "PENDING", "Xe chờ Admin duyệt."],
            ["CarStatus", "APPROVED", "Xe đã được duyệt, có thể hiển thị/cho thuê."],
            ["CarStatus", "RENTED", "Xe đang có lịch thuê chính thức hoặc đang được thuê."],
            ["CarStatus", "REJECTED", "Xe bị Admin từ chối duyệt."],
            ["CarStatus", "HIDDEN", "Xe bị ẩn theo trạng thái hệ thống."],
            ["BookingStatus", "REQUESTED", "Khách gửi yêu cầu thuê, chờ chủ xe duyệt."],
            ["BookingStatus", "OWNER_APPROVED", "Chủ xe đã duyệt, khách có thể thanh toán/tạo hợp đồng."],
            ["BookingStatus", "PAYMENT_PENDING", "Khách đã mở thanh toán, chờ kết quả."],
            ["BookingStatus", "PAID", "Đã thanh toán khoản cần trả trước."],
            ["BookingStatus", "IN_PROGRESS", "Xe đã bàn giao, chuyến thuê đang diễn ra."],
            ["BookingStatus", "COMPLETED", "Chuyến thuê hoàn tất."],
            ["BookingStatus", "CANCELLED", "Booking đã hủy."],
            ["BookingStatus", "REJECTED", "Chủ xe từ chối booking."],
            ["BookingStatus", "NO_SHOW", "Khách không đến nhận xe."],
            ["PaymentStatus", "PENDING", "Payment chờ xử lý."],
            ["PaymentStatus", "PAID", "Payment đã thanh toán."],
            ["PaymentStatus", "FAILED", "Payment thất bại."],
            ["PaymentStatus", "REFUNDED", "Payment đã hoàn tiền về mặt trạng thái."],
            ["CartStatus", "ACTIVE", "Giữ xe tạm thời còn hiệu lực."],
            ["CartStatus", "EXPIRED", "Giữ xe đã hết hạn."],
            ["CartStatus", "BOOKED", "Giữ xe đã chuyển thành booking."],
            ["CartStatus", "CANCELLED", "Giữ xe đã hủy."],
            ["ContractStatus", "DRAFT", "Hợp đồng nháp."],
            ["ContractStatus", "ACTIVE", "Hợp đồng đang hiệu lực."],
            ["ContractStatus", "COMPLETED", "Hợp đồng hoàn tất."],
            ["ContractStatus", "CANCELLED", "Hợp đồng đã hủy."],
        ],
        [1.4, 1.8, 3.9],
    )

    add_table(
        doc,
        ["Nhóm", "Giá trị"],
        [
            ["Role", "ADMIN, BUSINESS, USER"],
            ["OwnerType", "BUSINESS, USER"],
            ["BusinessType", "COMPANY, INDIVIDUAL"],
            ["CarType", "SUV, SEDAN, HATCHBACK, PICKUP, MPV, COUPE, CONVERTIBLE, ELECTRIC"],
            ["FuelType", "GASOLINE, DIESEL, ELECTRIC, HYBRID"],
            ["Transmission", "AUTOMATIC, MANUAL"],
            ["PaymentMethod", "CASH, MOMO, VNPAY"],
            ["PaymentOption", "DEPOSIT, FULL"],
            ["PaymentType", "DEPOSIT, FULL, REMAINING"],
            ["RentalMode", "DAILY, HOURLY"],
            ["PricingDateType", "WEEKDAY, WEEKEND, HOLIDAY"],
        ],
        [2.0, 4.8],
    )

    add_heading(doc, "7. Mô hình dữ liệu chính", 1)
    add_table(
        doc,
        ["Entity", "Dữ liệu chính", "Ghi chú nghiệp vụ"],
        [
            ["User", "name, email, password, phone, role, isVerified, isBlocked, isDeleted, otpCode, otpExpireAt", "Lưu tài khoản, phân quyền và trạng thái xác thực."],
            ["Business", "userId, businessName, businessType, phone, address, isApproved, isRejected, isDeleted", "Hồ sơ doanh nghiệp gắn với một User role BUSINESS."],
            ["Brand", "name, logo, isDeleted", "Danh mục thương hiệu xe."],
            ["Car", "ownerId, ownerType, brandId, name, type, licensePlate, pricing, pickup location, images, status, isHidden, isDeleted", "Xe có thể thuộc BUSINESS hoặc USER ký gửi."],
            ["Cart", "userId, carId, startDate, endDate, rentalMode, totalPrice, pricingSnapshot, expiredAt, status", "Giữ xe tạm thời trước khi tạo booking."],
            ["Booking", "userId, ownerId, ownerType, carId, startDate, endDate, totalPrice, paymentOption, paidAmount, remainingAmount, renterInfo, status", "Trung tâm của luồng thuê xe."],
            ["Payment", "bookingId, userId, amount, method, status, paymentType, paidAt, transactionCode", "Lưu từng lần thanh toán."],
            ["Contract", "bookingId, userId, carId, ownerId, renter snapshot, totalPrice, paidAmount, remainingAmount, paymentStatus, status", "Hợp đồng lấy dữ liệu snapshot từ booking."],
            ["HolidayCalendar", "name, date, isActive, isDeleted", "Danh sách ngày lễ phục vụ tính giá ngày lễ."],
        ],
        [1.3, 3.5, 2.8],
    )

    add_heading(doc, "8. Quy trình nghiệp vụ", 1)

    add_heading(doc, "8.1. Đăng ký tài khoản USER bằng OTP", 2)
    add_numbered(
        doc,
        [
            "Người dùng truy cập trang Đăng ký.",
            "Người dùng nhập email và bấm Gửi OTP.",
            "Hệ thống kiểm tra email, nếu email đã tồn tại và đã xác thực thì báo lỗi.",
            "Hệ thống sinh OTP 6 chữ số, thời hạn 5 phút, tạo hoặc cập nhật tài khoản tạm isVerified=false.",
            "Hệ thống gửi OTP qua email.",
            "Người dùng nhập họ tên, email, OTP, số điện thoại, mật khẩu, xác nhận mật khẩu và đồng ý điều khoản.",
            "Frontend kiểm tra dữ liệu bắt buộc, định dạng số điện thoại và độ mạnh mật khẩu.",
            "Hệ thống xác minh OTP: đúng mã, chưa hết hạn, tài khoản tồn tại.",
            "Nếu OTP hợp lệ, hệ thống chuyển isVerified=true và xóa otpCode/otpExpireAt.",
            "Hệ thống mã hóa mật khẩu bằng bcrypt, cập nhật tài khoản USER và chuyển người dùng sang trang đăng nhập.",
        ],
    )

    add_heading(doc, "8.2. Đăng nhập và phân quyền", 2)
    add_numbered(
        doc,
        [
            "Người dùng nhập email và mật khẩu hoặc đăng nhập bằng Google.",
            "Hệ thống kiểm tra tài khoản tồn tại, chưa bị xóa, chưa bị khóa và đã xác thực email.",
            "Với đăng nhập mật khẩu, hệ thống so sánh mật khẩu bằng bcrypt.",
            "Nếu hợp lệ, backend sinh JWT chứa thông tin người dùng và role.",
            "Frontend lưu token, user và role vào localStorage.",
            "Khi gọi API cần bảo vệ, frontend gửi token qua header Authorization: Bearer token.",
            "Backend kiểm tra token và role trước khi cho phép truy cập chức năng.",
        ],
    )

    add_heading(doc, "8.3. Quên mật khẩu", 2)
    add_numbered(
        doc,
        [
            "Người dùng nhập email tại trang quên mật khẩu.",
            "Hệ thống kiểm tra email hợp lệ và tài khoản tồn tại.",
            "Hệ thống sinh OTP đặt lại mật khẩu, mã hóa OTP và lưu thời hạn hiệu lực.",
            "Người dùng nhập OTP để xác minh.",
            "Nếu OTP đúng và còn hạn, hệ thống sinh resetToken tạm thời.",
            "Người dùng nhập mật khẩu mới và xác nhận mật khẩu.",
            "Hệ thống kiểm tra độ mạnh mật khẩu, xác nhận resetToken, mã hóa mật khẩu mới và xóa trạng thái đặt lại mật khẩu.",
        ],
    )

    add_heading(doc, "8.4. Admin cấp tài khoản BUSINESS", 2)
    add_numbered(
        doc,
        [
            "Admin đăng nhập và truy cập trang quản lý doanh nghiệp.",
            "Admin chọn Tạo doanh nghiệp và nhập tên doanh nghiệp, email, số điện thoại, địa chỉ, mô tả.",
            "Frontend kiểm tra dữ liệu bắt buộc, email hợp lệ và số điện thoại đúng định dạng.",
            "Admin bấm gửi OTP; backend kiểm tra email, tạo tài khoản tạm role BUSINESS, isVerified=false.",
            "Backend sinh OTP 6 chữ số, thời hạn 5 phút, lưu otpCode/otpExpireAt và gửi OTP qua email.",
            "Admin nhập OTP; backend kiểm tra tài khoản tạm, OTP đúng và chưa hết hạn.",
            "Nếu hợp lệ, backend chuyển isVerified=true và xóa OTP.",
            "Admin nhập mật khẩu và xác nhận mật khẩu cho tài khoản BUSINESS.",
            "Backend mã hóa mật khẩu bằng bcrypt, cập nhật User role BUSINESS, tạo Business với isApproved=true, isRejected=false.",
            "Hệ thống hiển thị tạo doanh nghiệp thành công và cập nhật lại danh sách doanh nghiệp.",
        ],
    )

    add_heading(doc, "8.5. Doanh nghiệp gửi yêu cầu trở thành BUSINESS", 2)
    add_numbered(
        doc,
        [
            "USER gửi yêu cầu doanh nghiệp với thông tin tên doanh nghiệp, loại hình, số điện thoại, địa chỉ và mô tả.",
            "Hệ thống tạo hồ sơ Business ở trạng thái chưa duyệt.",
            "Admin xem danh sách yêu cầu doanh nghiệp.",
            "Admin duyệt yêu cầu: user được chuyển role BUSINESS, Business isApproved=true.",
            "Admin từ chối yêu cầu: Business isRejected=true và lưu lý do từ chối.",
        ],
    )

    add_heading(doc, "8.6. Quản lý xe và duyệt xe", 2)
    add_numbered(
        doc,
        [
            "BUSINESS hoặc USER ký gửi tạo xe mới với thông tin thương hiệu, tên xe, loại xe, biển số, số ghế, nhiên liệu, hộp số, giá thuê, địa điểm nhận xe và hình ảnh.",
            "Backend kiểm tra dữ liệu bắt buộc, địa chỉ nhận xe, biển số không trùng và định dạng hợp lệ.",
            "Xe mới được lưu với trạng thái PENDING.",
            "Admin xem danh sách xe chờ duyệt, thông tin chủ xe, hình ảnh, giá thuê và vị trí nhận xe.",
            "Nếu duyệt, Admin chuyển status sang APPROVED và xe được hiển thị công khai.",
            "Nếu từ chối, Admin chuyển status sang REJECTED và lưu rejectReason.",
            "Chủ xe cập nhật xe: nếu thay đổi quan trọng thì xe về PENDING; nếu chỉ cập nhật vị trí qua luồng bản đồ thì lưu lịch sử vị trí và không cần duyệt lại.",
            "Chủ xe có thể ẩn/hiện xe hoặc xóa mềm xe nếu xe không có booking/hợp đồng đang hoạt động.",
        ],
    )

    add_heading(doc, "8.7. Tìm kiếm, xem chi tiết xe và bản đồ", 2)
    add_numbered(
        doc,
        [
            "Người thuê truy cập trang chủ hoặc trang tìm kiếm xe.",
            "Người thuê lọc theo địa điểm, thương hiệu, số ghế, khoảng giá, nhiên liệu, hộp số, loại xe, chế độ thuê và thời gian thuê.",
            "Hệ thống chỉ hiển thị xe APPROVED, không bị ẩn, không bị xóa mềm và không trùng lịch với booking/giữ xe hợp lệ.",
            "Người thuê xem chi tiết xe gồm hình ảnh, thông số, giá thuê, chính sách, vị trí nhận xe và trạng thái khả dụng.",
            "Bản đồ hiển thị vị trí nhận xe; người dùng có thể dùng vị trí hiện tại hoặc điểm bắt đầu tự chọn để tìm đường.",
        ],
    )

    add_heading(doc, "8.8. Giỏ xe và giữ xe tạm thời", 2)
    add_numbered(
        doc,
        [
            "Người dùng chọn xe và thời gian thuê.",
            "Hệ thống kiểm tra xe có thể đặt, không trùng booking và không trùng cart ACTIVE của người khác.",
            "Hệ thống tạo Cart với status ACTIVE, expiredAt và pricingSnapshot.",
            "Nếu cart hết hạn, hệ thống chuyển status sang EXPIRED.",
            "Người dùng có thể xóa xe khỏi giỏ, cart chuyển sang CANCELLED.",
            "Khi tạo booking từ giỏ, cart chuyển sang BOOKED.",
        ],
    )

    add_heading(doc, "8.9. Tạo booking trực tiếp hoặc từ giỏ", 2)
    add_numbered(
        doc,
        [
            "Người thuê chọn xe, ngày giờ thuê, phương thức thuê theo ngày/giờ và nhập thông tin người thuê.",
            "Thông tin người thuê gồm họ tên, số điện thoại, email, CCCD, ảnh CCCD, giấy phép lái xe và ghi chú.",
            "Hệ thống kiểm tra xe hợp lệ, không bị ẩn/xóa, trạng thái có thể đặt và không trùng lịch.",
            "Hệ thống tính giá thuê theo ngày thường, cuối tuần, ngày lễ hoặc theo giờ, sau đó lưu pricingSnapshot.",
            "Hệ thống tạo booking với status REQUESTED, chờ chủ xe duyệt.",
            "Nếu booking tạo từ giỏ, hệ thống đồng thời chuyển cart sang BOOKED.",
        ],
    )

    add_heading(doc, "8.10. Chủ xe duyệt hoặc từ chối booking", 2)
    add_numbered(
        doc,
        [
            "BUSINESS hoặc USER ký gửi xem danh sách booking thuộc xe của mình.",
            "Chủ xe xem chi tiết booking, thông tin người thuê và giấy tờ.",
            "Nếu đồng ý, chủ xe xác nhận booking; status chuyển sang OWNER_APPROVED.",
            "Nếu từ chối, chủ xe nhập lý do; status chuyển sang REJECTED.",
            "Sau khi được duyệt, người thuê có thể tạo hợp đồng và thanh toán.",
        ],
    )

    add_heading(doc, "8.11. Thanh toán", 2)
    add_numbered(
        doc,
        [
            "Người thuê chọn hình thức thanh toán: đặt cọc hoặc thanh toán toàn bộ.",
            "Người thuê chọn phương thức thanh toán CASH, MOMO hoặc VNPAY.",
            "Nếu chọn MoMo/VNPAY, hệ thống tạo payment PENDING, chuyển người dùng sang cổng thanh toán và xử lý return/callback.",
            "Khi thanh toán thành công, payment chuyển sang PAID; hệ thống đồng bộ paidAmount, remainingAmount và status booking.",
            "Nếu thanh toán cọc, hệ thống ghi nhận DEPOSIT và còn REMAINING cần thanh toán.",
            "Nếu thanh toán toàn bộ, paidAmount bằng totalPrice và remainingAmount bằng 0.",
            "Nếu còn tiền phần còn lại, người thuê có thể thanh toán tiếp trên hệ thống hoặc chủ xe xác nhận thu tiền mặt khi bàn giao.",
            "Chỉ payment PAID mới được tính vào số tiền đã thanh toán.",
        ],
    )

    add_heading(doc, "8.12. Hợp đồng thuê xe", 2)
    add_numbered(
        doc,
        [
            "Sau khi booking được chủ xe duyệt, người thuê có thể tạo hợp đồng.",
            "Hợp đồng lấy dữ liệu snapshot từ booking gồm người thuê, chủ xe, xe, thời gian thuê, địa điểm nhận xe, tổng tiền, tiền đã thanh toán và còn lại.",
            "Hợp đồng có trạng thái ACTIVE khi booking đang hiệu lực.",
            "Khi booking COMPLETED, CANCELLED, REJECTED hoặc NO_SHOW, trạng thái hợp đồng được đồng bộ tương ứng.",
            "Người thuê và chủ xe có thể xem danh sách hợp đồng và chi tiết hợp đồng liên quan đến mình.",
        ],
    )

    add_heading(doc, "8.13. Bàn giao, hoàn tất và no-show", 2)
    add_numbered(
        doc,
        [
            "Sau khi booking được duyệt và đáp ứng điều kiện thanh toán trước, chủ xe có thể bàn giao xe.",
            "Khi bàn giao, booking chuyển sang IN_PROGRESS và xe chuyển sang RENTED.",
            "Nếu còn tiền phần còn lại và người thuê trả trực tiếp, chủ xe xác nhận đã thu tiền mặt; hệ thống tạo payment REMAINING với method CASH và status PAID.",
            "Khi người thuê trả xe, chủ xe hoàn tất booking.",
            "Hệ thống chỉ cho hoàn tất khi booking đang IN_PROGRESS và remainingAmount bằng 0.",
            "Khi hoàn tất, booking chuyển sang COMPLETED, hợp đồng chuyển COMPLETED và xe được giải phóng về APPROVED nếu không còn booking active.",
            "Nếu khách không đến nhận xe đúng lịch, chủ xe có thể xử lý no-show; booking chuyển sang NO_SHOW và xe được mở lại.",
        ],
    )

    add_heading(doc, "8.14. Ngày lễ và tính giá", 2)
    add_numbered(
        doc,
        [
            "Admin quản lý danh sách ngày lễ: thêm, sửa, xóa mềm và bật/tắt.",
            "Xe hỗ trợ giá ngày thường, cuối tuần, ngày lễ và giá theo giờ nếu bật cho thuê theo giờ.",
            "Khi tính giá, hệ thống ưu tiên ngày lễ, sau đó cuối tuần, sau đó ngày thường.",
            "Khi tạo cart/booking, hệ thống lưu pricingSnapshot để giữ nguyên giá tại thời điểm đặt.",
        ],
    )

    add_heading(doc, "8.15. Thông báo và dashboard", 2)
    add_bullets(
        doc,
        [
            "Dashboard ADMIN tổng hợp số lượng người dùng, doanh nghiệp, xe, booking, doanh thu, xe chờ duyệt và các thống kê phân loại.",
            "Dashboard BUSINESS tổng hợp số xe, booking, trạng thái xe và doanh thu của doanh nghiệp.",
            "Khu vực USER ký gửi có dashboard riêng cho xe ký gửi, booking và thanh toán.",
            "USER thuê xe có các trang theo dõi booking, hợp đồng và thanh toán; chưa có dashboard USER thường riêng.",
            "Notification summary hiển thị số lượng việc cần xử lý theo từng role như xe chờ duyệt, booking mới, xe bị từ chối, booking cần bàn giao hoặc hợp đồng mới.",
        ],
    )

    add_heading(doc, "9. Quy tắc nghiệp vụ quan trọng", 1)
    add_bullets(
        doc,
        [
            "Chỉ xe đã APPROVED và không bị ẩn/xóa mềm mới được hiển thị cho người thuê.",
            "Không cho người dùng đặt xe của chính mình.",
            "Không cho đặt trùng thời gian với booking hoặc cart đang giữ hợp lệ.",
            "Booking mới luôn ở trạng thái REQUESTED và phải chờ chủ xe duyệt trước khi thanh toán.",
            "Xe đang có booking đã thanh toán hoặc đang thuê có thể được đánh dấu RENTED.",
            "Xe RENTED không được Admin thay đổi trạng thái duyệt.",
            "Chủ xe không được xóa xe nếu xe có booking/hợp đồng đang hoạt động.",
            "Admin khóa doanh nghiệp hoặc user sở hữu xe thì các xe liên quan có thể bị ẩn khỏi hệ thống.",
            "Payment PENDING, FAILED hoặc REFUNDED không được cộng vào paidAmount.",
            "Booking không được hoàn tất nếu còn remainingAmount.",
            "Mail lỗi không được làm hỏng nghiệp vụ chính trong các luồng thông báo bất đồng bộ.",
            "Dữ liệu địa chỉ và vị trí nhận xe được lưu snapshot trong booking/contract để tránh thay đổi lịch sử khi xe cập nhật địa chỉ sau này.",
        ],
    )

    add_heading(doc, "10. Kiểm tra dữ liệu đầu vào", 1)
    add_table(
        doc,
        ["Nhóm dữ liệu", "Kiểm tra"],
        [
            ["Email", "Bắt buộc ở các luồng tài khoản, đúng định dạng, không trùng với tài khoản đã xác thực."],
            ["Mật khẩu USER", "Tối thiểu 8 ký tự, không chứa khoảng trắng, có chữ hoa, chữ thường và số."],
            ["Mật khẩu BUSINESS do Admin cấp", "Frontend kiểm tra tối thiểu 6 ký tự và xác nhận mật khẩu trùng khớp."],
            ["Số điện thoại", "10 chữ số và bắt đầu bằng 0."],
            ["Biển số xe", "Được chuẩn hóa và kiểm tra không trùng."],
            ["Địa chỉ nhận xe", "Bắt buộc khi đăng/cập nhật xe."],
            ["Ngày thuê", "Ngày trả phải sau ngày nhận; thời gian thuê phải phù hợp daily/hourly."],
            ["Thông tin người thuê", "Bắt buộc các thông tin cá nhân và giấy tờ cần thiết khi tạo booking."],
            ["Thanh toán", "Chỉ cho tạo payment khi booking ở trạng thái cho phép và paymentType hợp lệ."],
        ],
        [2.0, 4.8],
    )

    add_heading(doc, "11. Tích hợp ngoài", 1)
    add_table(
        doc,
        ["Tích hợp", "Mục đích", "Ghi chú"],
        [
            ["Google OAuth", "Đăng nhập nhanh bằng tài khoản Google.", "Tài khoản Google có thể tự xác thực email."],
            ["SMTP/Gmail", "Gửi OTP, quên mật khẩu và email thông báo.", "Sử dụng Nodemailer."],
            ["MoMo", "Tạo link thanh toán và xử lý IPN/return.", "Hỗ trợ thanh toán online."],
            ["VNPAY", "Tạo URL thanh toán và xử lý return.", "Hỗ trợ thanh toán online."],
            ["Map/Geocode/Route", "Tìm tọa độ địa chỉ và tìm đường.", "Frontend dùng Leaflet/react-leaflet để hiển thị."],
        ],
        [1.5, 2.3, 3.0],
    )

    add_heading(doc, "12. Tổng quan API chính", 1)
    add_table(
        doc,
        ["Module", "Endpoint chính", "Chức năng"],
        [
            ["Auth", "/auth/send-otp, /auth/verify-otp, /auth/register, /auth/login, /auth/google-login", "Tài khoản, OTP, đăng nhập và Google login."],
            ["Admin", "/admin/users, /admin/business/*", "Quản lý user, doanh nghiệp và cấp tài khoản BUSINESS."],
            ["Brand", "/brand/createBrand, /brand/getAllBrand, /brand/updateBrand/:id, /brand/deleteBrand/:id", "Quản lý thương hiệu xe."],
            ["Business", "/business/dashboard, /business/profile, /business/requestBusiness", "Hồ sơ, dashboard và yêu cầu doanh nghiệp."],
            ["Car", "/cars/createCar, /cars/search, /cars/getOneCar/:id, /cars/updateCar/:id, /cars/approveCar/:id", "Quản lý xe, tìm kiếm và duyệt xe."],
            ["Cart", "/cart/addToCart, /cart/getMyCart, /cart/removeFromCart/:id", "Giữ xe tạm thời."],
            ["Booking", "/bookings/createBooking, /bookings/bookingFromCart/:cartId, /bookings/confirmBooking/:id, /bookings/handoverBooking/:id", "Booking, duyệt, bàn giao, hoàn tất và no-show."],
            ["Payment", "/payments/createPayment, /payments/momo/create, /payments/vnpay/create, /payments/updatePaymentStatus/:id", "Thanh toán và lịch sử thanh toán."],
            ["Contract", "/contracts/create, /contracts/my-contracts, /contracts/owner/my-contracts, /contracts/:id", "Hợp đồng thuê xe."],
            ["Holiday", "/holidays/public, /admin/holidays", "Ngày lễ và giá ngày lễ."],
            ["Maps", "/maps/geocode, /maps/route", "Geocode và tìm đường."],
            ["Owner", "/owner/cars/map, /owner/cars/:id/location", "Quản lý vị trí xe của chủ xe."],
            ["Notification", "/notifications/summary", "Tổng hợp thông báo theo vai trò."],
            ["Dashboard", "/dashboard/admin, /dashboard/business", "Thống kê cho Admin và Business."],
        ],
        [1.3, 3.2, 2.5],
    )

    add_heading(doc, "13. Giới hạn hiện tại và hướng phát triển", 1)
    add_table(
        doc,
        ["Nội dung", "Trạng thái hiện tại", "Khuyến nghị viết trong báo cáo"],
        [
            ["Trạng thái Maintenance của xe", "Chưa có trong enum/model.", "Không mô tả là chức năng đã có; có thể đưa vào hướng phát triển."],
            ["Available là trạng thái chính của xe", "Không lưu trong CarStatus; xe sẵn sàng tương ứng APPROVED, còn AVAILABLE chỉ là rentalAvailability phụ.", "Dùng APPROVED cho trạng thái duyệt/sẵn sàng."],
            ["BANKING/manual hoặc DIRECT", "Chưa có PaymentMethod tương ứng.", "Chỉ ghi CASH, MOMO, VNPAY."],
            ["Phí phát sinh sau chuyến thuê", "Chưa có model/route/paymentType EXTRA_CHARGE ở backend.", "Ghi là hướng phát triển nếu cần."],
            ["Hoàn tiền đầy đủ", "Có status REFUNDED nhưng chưa có luồng hoàn tiền hoàn chỉnh.", "Không mô tả là nghiệp vụ hoàn tiền tự động đã hoàn thiện."],
            ["Dashboard USER thường", "Có trang booking/payment/contract; dashboard riêng chủ yếu cho USER ký gửi.", "Viết USER có trang quản lý, USER ký gửi có dashboard."],
            ["Cập nhật hồ sơ USER thường", "Chưa thấy endpoint cập nhật profile USER thường; BUSINESS có /business/profile.", "Viết hẹp theo chức năng hiện có."],
        ],
        [2.0, 2.8, 2.2],
    )

    add_heading(doc, "14. Kết luận", 1)
    doc.add_paragraph(
        "BQDrive hiện là hệ thống đặt thuê xe trực tuyến có đầy đủ các module cốt lõi: tài khoản, phân quyền, quản lý xe, duyệt xe, tìm kiếm, booking, thanh toán, hợp đồng, bản đồ, ngày lễ, dashboard và thông báo. Hệ thống đã phân tách rõ vai trò ADMIN, BUSINESS, USER thuê xe và USER ký gửi. Các nghiệp vụ đang được tổ chức theo RESTful API giữa frontend React và backend ExpressJS, dữ liệu lưu trong MongoDB thông qua Mongoose."
    )

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
